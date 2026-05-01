from __future__ import annotations

from pathlib import Path

from anon_tool.types import InputLine


def read_docx_lines(path: Path) -> list[InputLine]:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "Missing dependency 'python-docx'. Install project dependencies before running DOCX input."
        ) from exc

    document = Document(str(path))
    lines: list[InputLine] = []

    def add_text(text: str) -> None:
        lines.append(InputLine(page=1, line_no=len(lines) + 1, text=text))

    for paragraph in document.paragraphs:
        add_text(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            add_text(" | ".join(cells))

    return lines
