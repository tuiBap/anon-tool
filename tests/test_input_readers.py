from __future__ import annotations

import importlib.util
from pathlib import Path

import pytest

from anon_tool.cli import _read_input, _resolve_input_type, _write_output
from anon_tool.types import InputLine


def test_resolves_supported_input_types() -> None:
    assert _resolve_input_type(Path("case.pdf"), "auto") == "pdf"
    assert _resolve_input_type(Path("case.txt"), "auto") == "txt"
    assert _resolve_input_type(Path("case.docx"), "auto") == "docx"


def test_rejects_unsupported_input_type() -> None:
    with pytest.raises(ValueError, match="Unable to infer input type"):
        _resolve_input_type(Path("case.doc"), "auto")


def test_writes_supported_output_formats(tmp_path: Path) -> None:
    lines = [InputLine(page=1, line_no=1, text="redacted")]
    markdown_path = tmp_path / "case.md"
    text_path = tmp_path / "case.txt"
    pdf_path = tmp_path / "case.pdf"

    _write_output(markdown_path, lines, "markdown")
    _write_output(text_path, lines, "text")
    _write_output(pdf_path, lines, "pdf")

    assert markdown_path.read_text(encoding="utf-8") == (
        "# Sanitized Case Record\n\n## Case Content\n\n### Source Page 1\n\nredacted\n"
    )
    assert text_path.read_text(encoding="utf-8") == "=== Source Page 1 ===\nredacted\n"
    assert pdf_path.read_bytes().startswith(b"%PDF")


@pytest.mark.skipif(importlib.util.find_spec("docx") is None, reason="python-docx not installed")
def test_reads_docx_lines(tmp_path: Path) -> None:
    from docx import Document

    input_path = tmp_path / "case.docx"
    document = Document()
    document.add_paragraph("Created By David Bush")
    document.add_paragraph("Phone: 847-267-9330")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Case"
    table.cell(0, 1).text = "12345678"
    document.save(input_path)

    lines = _read_input(input_path, "docx")

    assert [line.text for line in lines] == [
        "Created By David Bush",
        "Phone: 847-267-9330",
        "Case | 12345678",
    ]
